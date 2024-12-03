from docx import Document
from docx.shared import RGBColor, Pt
from bs4 import BeautifulSoup
import requests
import shutil


def create_pdf(quiz):
    document = Document()

    par = document.add_paragraph().add_run("КвоккаКвиз Офлайн")
    font = par.font
    font.name = "Times New Roman"
    font.size = Pt(25)
    font.color.rgb = RGBColor(000, 000, 000)
    font.bold = True

    par = document.add_paragraph().add_run(f"Квиз №{quiz['id']} - {quiz['title']}")
    font = par.font
    font.name = "Times New Roman"
    font.size = Pt(15)
    font.color.rgb = RGBColor(000, 000, 000)

    par = document.add_paragraph().add_run("Фамилия: ___________ Имя: ___________ Класс: ___ Буква: ___")
    font = par.font
    font.name = "Times New Roman"
    font.size = Pt(15)
    font.color.rgb = RGBColor(000, 000, 000)

    counter = 1

    for t in quiz['test']:
        par = document.add_paragraph().add_run(f"№ {counter}/{len(quiz['test'])}")
        font = par.font
        font.name = "Times New Roman"
        font.size = Pt(20)
        font.color.rgb = RGBColor(000, 000, 000)
        font.bold = True

        if t['image'] != "None":
            response = requests.get(t['image'], stream=True)
            if response.status_code == 200:
                with open('image.jpg', 'wb') as file:
                    response.raw.decode_content = True
                    shutil.copyfileobj(response.raw, file)
                    document.add_picture('image.jpg')
                print('Фото успешно скачано')
            else:
                print('Ошибка при загрузке фото')

        cleantext = BeautifulSoup(t['text'], "lxml").text

        par = document.add_paragraph().add_run(cleantext)
        font = par.font
        font.name = "Times New Roman"
        font.size = Pt(15)
        font.color.rgb = RGBColor(000, 000, 000)

        if t['type'] == 'text':
            par = document.add_paragraph().add_run("Ответ: _______________")
            font = par.font
            font.name = "Times New Roman"
            font.size = Pt(15)
            font.color.rgb = RGBColor(000, 000, 000)
        elif t['type'] == 'choice':
            answer_counter = 1
            for a in t['answer']:
                par = document.add_paragraph().add_run(f"{answer_counter}) {a['answer']}")
                font = par.font
                font.name = "Times New Roman"
                font.size = Pt(14)
                font.color.rgb = RGBColor(000, 000, 000)
                answer_counter += 1
            par = document.add_paragraph().add_run("Ответ: _____")
            font = par.font
            font.name = "Times New Roman"
            font.size = Pt(15)
            font.color.rgb = RGBColor(000, 000, 000)

        counter += 1

    document.save('static/tasks.docx')